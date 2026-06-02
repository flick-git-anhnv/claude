"""
md_to_docx_kztek.py — Chuyển đổi Markdown → DOCX + PDF + HTML theo chuẩn thương hiệu KZTEK

Yêu cầu:
    pip install python-docx Pillow

Chuyển đổi sang PDF (chọn 1 trong các cách):
    pip install docx2pdf          # Windows / macOS (cần MS Word)
    sudo apt install libreoffice  # Linux (dùng: soffice --headless)

Cách dùng:
    # DOCX + PDF (mặc định)
    python scripts/md_to_docx_kztek.py docs/prd/PRD-iled-parking.md

    # Chỉ xuất DOCX, bỏ qua PDF
    python scripts/md_to_docx_kztek.py docs/prd/PRD-iled-parking.md --no-pdf

    # Xuất thêm HTML
    python scripts/md_to_docx_kztek.py docs/prd/PRD-iled-parking.md --html

    # Chỉ xuất HTML (bỏ DOCX/PDF)
    python scripts/md_to_docx_kztek.py docs/prd/PRD-iled-parking.md --html --no-docx --no-pdf

    # HTML tự chứa (logo + ảnh nhúng base64, không cần file ngoài)
    python scripts/md_to_docx_kztek.py docs/prd/PRD-x.md --html --embed-assets

    # Chỉ định thư mục output
    python scripts/md_to_docx_kztek.py docs/prd/PRD-iled-parking.md --output-dir exports/

    # Toàn bộ file *.md trong thư mục
    python scripts/md_to_docx_kztek.py docs/prd/ --batch

    # Nhiều file cụ thể
    python scripts/md_to_docx_kztek.py docs/prd/PRD-x.md docs/user-stories/US-001.md
"""

import argparse
import base64
import os
import re
import subprocess
import sys
import webbrowser
from pathlib import Path

# ── python-docx ─────────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt, RGBColor
except ImportError:
    sys.exit("[KZTEK] Thiếu thư viện: pip install python-docx")

# ── KZTEK Brand Colors ───────────────────────────────────────────────────────
NAVY_DARK   = RGBColor(0x25, 0x1C, 0x53)   # #251C53 — heading chính   (60–70%)
NAVY_LIGHT  = RGBColor(0x4A, 0x3F, 0x8C)   # #4A3F8C — sub-heading, link
NAVY_PALE   = RGBColor(0xB8, 0xB3, 0xD6)   # #B8B3D6 — fill bảng xen kẽ
ORANGE      = RGBColor(0xF0, 0x59, 0x22)   # #F05922 — accent, CTA     (15–20%)
ORANGE_PALE = RGBColor(0xFF, 0xAA, 0x80)   # #FFAA80 — highlight nhẹ
GRAY        = RGBColor(0xCB, 0xCB, 0xCB)   # #CBCBCB — border, divider
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)   # #FFFFFF — nền chính

# Hex string cho XML attribute (không có #)
NAVY_DARK_HEX = "251C53"
NAVY_PALE_HEX = "B8B3D6"
ORANGE_HEX    = "F05922"
WHITE_HEX     = "FFFFFF"
GRAY_HEX      = "CBCBCB"

# Logo — tìm theo thứ tự ưu tiên
LOGO_CANDIDATES = [
    "Kztek_Logo.png",
    "Kztek_Logo.jpg",
    ".claude/commands/Kztek_Logo.jpg",
    ".claude/commands/Kztek_Logo.png",
]


# ── Helpers ──────────────────────────────────────────────────────────────────

def find_logo() -> str | None:
    for path in LOGO_CANDIDATES:
        if os.path.exists(path):
            return path
    return None


def logo_to_base64(logo_path: str) -> tuple[str, str]:
    """Trả về (data_uri, mime_type) để nhúng thẳng vào HTML."""
    ext = Path(logo_path).suffix.lower()
    mime = "image/png" if ext == ".png" else "image/jpeg"
    with open(logo_path, "rb") as f:
        data = base64.b64encode(f.read()).decode("ascii")
    return f"data:{mime};base64,{data}", mime


def image_to_base64(img_path: str) -> str | None:
    """Chuyển ảnh thành data URI để nhúng vào HTML."""
    try:
        ext = Path(img_path).suffix.lower()
        mime_map = {".png": "image/png", ".jpg": "image/jpeg",
                    ".jpeg": "image/jpeg", ".gif": "image/gif",
                    ".svg": "image/svg+xml", ".webp": "image/webp"}
        mime = mime_map.get(ext, "image/png")
        with open(img_path, "rb") as f:
            data = base64.b64encode(f.read()).decode("ascii")
        return f"data:{mime};base64,{data}"
    except Exception:
        return None


def set_cell_bg(cell, hex_color: str):
    """Đặt màu nền cho ô bảng."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, border_color: str = GRAY_HEX):
    """Đặt border mỏng cho ô bảng."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), border_color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_horizontal_rule(doc: Document):
    """Thêm đường kẻ ngang màu cam (thay thế ---)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), ORANGE_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)


def apply_inline_formatting(run, text: str):
    """Áp dụng bold/italic từ markdown inline."""
    # Đơn giản: nếu text bắt đầu/kết thúc bằng ** → bold; * → italic
    if text.startswith("**") and text.endswith("**"):
        run.text = text[2:-2]
        run.bold = True
    elif text.startswith("*") and text.endswith("*"):
        run.text = text[1:-1]
        run.italic = True
    else:
        run.text = text


def parse_inline(paragraph, text: str, base_color: RGBColor = None):
    """Parse bold/italic/code trong 1 dòng và thêm vào paragraph."""
    # Pattern: **bold**, *italic*, `code`
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if base_color:
            run.font.color.rgb = base_color
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = ORANGE
        else:
            run.text = part


# ── DOCX Builder ─────────────────────────────────────────────────────────────

def build_doc_header(doc: Document, logo_path: str | None):
    """Tạo header: Logo KZTEK + tên công ty + đường kẻ màu cam."""
    header = doc.sections[0].header
    # Xóa đoạn mặc định
    for para in header.paragraphs:
        para.clear()

    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if logo_path:
        run_logo = hp.add_run()
        try:
            run_logo.add_picture(logo_path, height=Inches(0.38))
        except Exception:
            pass  # Không crash nếu logo lỗi

    run_name = hp.add_run("   CÔNG TY CỔ PHẦN ĐẦU TƯ VÀ PHÁT TRIỂN KZTEK")
    run_name.font.size = Pt(8)
    run_name.font.color.rgb = NAVY_DARK
    run_name.font.bold = True

    # Đường kẻ cam dưới header
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), ORANGE_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_doc_footer(doc: Document):
    """Footer: số trang + website KZTEK."""
    footer = doc.sections[0].footer
    for para in footer.paragraphs:
        para.clear()

    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_left = fp.add_run("kztek.net  |  sales@kztek.net  |  0988 637 099")
    run_left.font.size = Pt(8)
    run_left.font.color.rgb = NAVY_LIGHT

    fp.add_run("     Trang ")
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run_page = fp.add_run()
    run_page._r.append(fldChar1)
    run_page._r.append(instrText)
    run_page._r.append(fldChar2)
    run_page.font.size = Pt(8)
    run_page.font.color.rgb = NAVY_LIGHT


def convert_md_to_docx(md_path: Path, output_path: Path, logo_path: str | None):
    """Đọc file Markdown và tạo DOCX theo brand KZTEK."""
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)

    build_doc_header(doc, logo_path)
    build_doc_footer(doc)

    i = 0
    in_code_block = False
    code_lines: list[str] = []
    in_table = False
    table_rows: list[list[str]] = []
    in_ordered_list = False

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False
            return

        # Lọc bỏ dòng phân cách |---|---|
        data_rows = [r for r in table_rows if not re.match(r"^\|[\s\-|:]+\|$", r[0] if r else "")]

        if not data_rows:
            in_table = False
            table_rows = []
            return

        col_count = max(len(r) for r in data_rows)
        tbl = doc.add_table(rows=len(data_rows), cols=col_count)
        tbl.style = "Table Grid"

        for ri, row_cells in enumerate(data_rows):
            for ci in range(col_count):
                cell = tbl.cell(ri, ci)
                cell_text = row_cells[ci].strip() if ci < len(row_cells) else ""

                # Header row: nền Navy, chữ trắng, bold
                if ri == 0:
                    set_cell_bg(cell, NAVY_DARK_HEX)
                    p = cell.paragraphs[0]
                    run = p.add_run(cell_text)
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(10)
                    set_cell_border(cell, NAVY_DARK_HEX)
                else:
                    # Dòng xen kẽ: trắng / navy pale
                    if ri % 2 == 0:
                        set_cell_bg(cell, NAVY_PALE_HEX)
                    p = cell.paragraphs[0]
                    parse_inline(p, cell_text)
                    for run in p.runs:
                        run.font.size = Pt(10)
                    set_cell_border(cell)

        doc.add_paragraph()
        in_table = False
        table_rows = []

    def flush_code(lines_buf: list[str]):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        # Background xám nhạt qua shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F4F4F8")
        pPr.append(shd)
        run = p.add_run("\n".join(lines_buf))
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = NAVY_DARK

    def force_restart_numbering(para):
        """Buộc danh sách có thứ tự bắt đầu lại từ 1 (tạo numId mới với startOverride=1)."""
        try:
            numbering_part = doc.part.numbering_part
            if numbering_part is None:
                return
            numbering = numbering_part._element
            nums = numbering.findall(qn("w:num"))
            if not nums:
                return
            pPr = para._p.find(qn("w:pPr"))
            if pPr is None:
                return
            numPr = pPr.find(qn("w:numPr"))
            if numPr is None:
                return
            numId_el = numPr.find(qn("w:numId"))
            if numId_el is None:
                return
            cur_id = numId_el.get(qn("w:val"))
            abstract_id = None
            for num in nums:
                if num.get(qn("w:numId")) == cur_id:
                    ab = num.find(qn("w:abstractNumId"))
                    if ab is not None:
                        abstract_id = ab.get(qn("w:val"))
                    break
            if abstract_id is None:
                return
            new_id = str(max(int(n.get(qn("w:numId"))) for n in nums) + 1)
            new_num = OxmlElement("w:num")
            new_num.set(qn("w:numId"), new_id)
            ab_ref = OxmlElement("w:abstractNumId")
            ab_ref.set(qn("w:val"), abstract_id)
            new_num.append(ab_ref)
            lvl = OxmlElement("w:lvlOverride")
            lvl.set(qn("w:ilvl"), "0")
            start = OxmlElement("w:startOverride")
            start.set(qn("w:val"), "1")
            lvl.append(start)
            new_num.append(lvl)
            numbering.append(new_num)
            numId_el.set(qn("w:val"), new_id)
        except Exception:
            pass

    while i < len(lines):
        line = lines[i]

        # ── Code block ────────────────────────────────────────────────────────
        if line.startswith("```"):
            if not in_code_block:
                if in_table:
                    flush_table()
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                flush_code(code_lines)
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Table ─────────────────────────────────────────────────────────────
        if line.startswith("|"):
            if not in_table:
                in_table = True
                table_rows = []
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # ── Heading H1 ────────────────────────────────────────────────────────
        if line.startswith("# ") and not line.startswith("## "):
            in_ordered_list = False
            heading = doc.add_heading(line[2:].strip(), level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in heading.runs:
                run.font.color.rgb = NAVY_DARK
                run.font.size = Pt(18)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after  = Pt(6)
            i += 1
            continue

        # ── Heading H2 ────────────────────────────────────────────────────────
        if line.startswith("## ") and not line.startswith("### "):
            in_ordered_list = False
            heading = doc.add_heading(line[3:].strip(), level=2)
            for run in heading.runs:
                run.font.color.rgb = ORANGE
                run.font.size = Pt(14)
            heading.paragraph_format.space_before = Pt(10)
            heading.paragraph_format.space_after  = Pt(4)
            i += 1
            continue

        # ── Heading H3 ────────────────────────────────────────────────────────
        if line.startswith("### ") and not line.startswith("#### "):
            in_ordered_list = False
            heading = doc.add_heading(line[4:].strip(), level=3)
            for run in heading.runs:
                run.font.color.rgb = NAVY_LIGHT
                run.font.size = Pt(12)
            heading.paragraph_format.space_before = Pt(8)
            heading.paragraph_format.space_after  = Pt(3)
            i += 1
            continue

        # ── Heading H4 ────────────────────────────────────────────────────────
        if line.startswith("#### "):
            in_ordered_list = False
            heading = doc.add_heading(line[5:].strip(), level=4)
            for run in heading.runs:
                run.font.color.rgb = NAVY_DARK
                run.font.size = Pt(11)
            i += 1
            continue

        # ── Horizontal rule ───────────────────────────────────────────────────
        if re.match(r"^[-*_]{3,}$", line.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Blockquote ────────────────────────────────────────────────────────
        if line.startswith("> "):
            content = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(1.0)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            # Đường kẻ trái màu cam
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "12")
            left.set(qn("w:space"), "10")
            left.set(qn("w:color"), ORANGE_HEX)
            pBdr.append(left)
            pPr.append(pBdr)
            parse_inline(p, content, NAVY_LIGHT)
            i += 1
            continue

        # ── Unordered list ────────────────────────────────────────────────────
        if re.match(r"^(\s*)[-*+] ", line):
            m = re.match(r"^(\s*)[-*+] (.*)", line)
            indent_level = len(m.group(1)) // 2
            content = m.group(2).strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent  = Cm(0.5 + indent_level * 0.5)
            p.paragraph_format.space_after  = Pt(1)
            parse_inline(p, content)
            for run in p.runs:
                run.font.size = Pt(10.5)
            i += 1
            continue

        # ── Ordered list ──────────────────────────────────────────────────────
        if re.match(r"^(\s*)\d+\. ", line):
            m = re.match(r"^(\s*)(\d+)\. (.*)", line)
            indent_level = len(m.group(1)) // 2
            num     = m.group(2)
            content = m.group(3).strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent       = Cm(1.0 + indent_level * 0.5)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.paragraph_format.space_after       = Pt(1)
            run_num = p.add_run(f"{num}.\t")
            run_num.font.size = Pt(10.5)
            parse_inline(p, content)
            for run in p.runs:
                run.font.size = Pt(10.5)
            in_ordered_list = True
            i += 1
            continue

        # ── Checkbox list  - [ ] / - [x] ─────────────────────────────────────
        if re.match(r"^(\s*)[-*] \[[ xX]\] ", line):
            m = re.match(r"^(\s*)[-*] \[([ xX])\] (.*)", line)
            checked = m.group(2).lower() == "x"
            content  = m.group(3).strip()
            symbol = "☑" if checked else "☐"
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(1)
            run_sym = p.add_run(f"{symbol}  ")
            run_sym.font.color.rgb = ORANGE if checked else GRAY
            run_sym.font.size = Pt(11)
            parse_inline(p, content)
            for run in p.runs[1:]:
                run.font.size = Pt(10.5)
                if checked:
                    run.font.color.rgb = NAVY_LIGHT
            i += 1
            continue

        # ── Image ─────────────────────────────────────────────────────────────
        img_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if img_match:
            alt       = img_match.group(1)
            img_path  = img_match.group(2)

            # Resolve tương đối so với thư mục chứa file .md
            resolved  = md_path.parent / img_path
            abs_path  = str(resolved.resolve()) if resolved.exists() else (
                        img_path if os.path.exists(img_path) else None)

            if abs_path:
                try:
                    # Đoạn trước ảnh: thêm khoảng trắng nhỏ
                    doc.add_paragraph().paragraph_format.space_before = Pt(4)

                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after  = Pt(0)

                    # Thêm border nhẹ quanh đoạn chứa ảnh
                    pPr = p._p.get_or_add_pPr()
                    pBdr = OxmlElement("w:pBdr")
                    for side in ("top", "left", "bottom", "right"):
                        el = OxmlElement(f"w:{side}")
                        el.set(qn("w:val"), "single")
                        el.set(qn("w:sz"), "4")
                        el.set(qn("w:space"), "4")
                        el.set(qn("w:color"), GRAY_HEX)
                        pBdr.append(el)
                    pPr.append(pBdr)

                    run = p.add_run()
                    run.add_picture(abs_path, width=Inches(5.8))

                    # Caption: dòng nghiêng ngay dưới ảnh, màu navy nhạt
                    if alt:
                        caption_p = doc.add_paragraph()
                        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_p.paragraph_format.space_before = Pt(2)
                        caption_p.paragraph_format.space_after  = Pt(8)
                        r = caption_p.add_run(alt)
                        r.font.size    = Pt(9)
                        r.font.italic  = True
                        r.font.color.rgb = NAVY_LIGHT
                    else:
                        doc.add_paragraph().paragraph_format.space_after = Pt(8)

                except Exception as ex:
                    p = doc.add_paragraph(f"[Không thể chèn ảnh: {img_path} — {ex}]")
                    for r in p.runs:
                        r.font.color.rgb = ORANGE
            else:
                # Ảnh không tìm thấy — cảnh báo rõ trong tài liệu
                p = doc.add_paragraph(f"⚠ ẢNH CHƯA CÓ: {img_path}")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.color.rgb = ORANGE
                    r.font.bold      = True
                    r.font.size      = Pt(10)
            i += 1
            continue

        # ── Dòng trắng ────────────────────────────────────────────────────────
        if line.strip() == "":
            in_ordered_list = False
            i += 1
            continue

        # ── Page break hint: <!-- pagebreak --> ───────────────────────────────
        if "<!-- pagebreak -->" in line.lower():
            doc.add_page_break()
            i += 1
            continue

        # ── Đoạn văn bản thường ───────────────────────────────────────────────
        in_ordered_list = False
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        parse_inline(p, line.strip())
        for run in p.runs:
            if not run.font.size:
                run.font.size = Pt(10.5)
        i += 1

    # Flush nếu kết thúc file đang trong table/code
    if in_table:
        flush_table()
    if in_code_block and code_lines:
        flush_code(code_lines)

    doc.save(str(output_path))
    return output_path


# ── PDF Conversion ───────────────────────────────────────────────────────────

def convert_docx_to_pdf(docx_path: Path) -> Path | None:
    """Chuyển DOCX → PDF. Thử nhiều phương án theo thứ tự."""
    pdf_path = docx_path.with_suffix(".pdf")

    # Phương án 1: docx2pdf (Windows / macOS, cần MS Word)
    try:
        from docx2pdf import convert as d2p_convert
        d2p_convert(str(docx_path), str(pdf_path))
        if pdf_path.exists():
            return pdf_path
    except ImportError:
        pass
    except Exception as e:
        print(f"  [docx2pdf] Lỗi: {e}")

    # Phương án 2: LibreOffice (Linux / cross-platform)
    for soffice_cmd in ("soffice", "libreoffice"):
        try:
            result = subprocess.run(
                [soffice_cmd, "--headless", "--convert-to", "pdf",
                 str(docx_path), "--outdir", str(docx_path.parent)],
                capture_output=True, timeout=60
            )
            if result.returncode == 0 and pdf_path.exists():
                return pdf_path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    # Phương án 3: pypandoc (cần pandoc + pdflatex)
    try:
        import pypandoc
        pypandoc.convert_file(
            str(docx_path), "pdf",
            outputfile=str(pdf_path),
            extra_args=["--pdf-engine=xelatex"]
        )
        if pdf_path.exists():
            return pdf_path
    except ImportError:
        pass
    except Exception as e:
        print(f"  [pypandoc] Lỗi: {e}")

    print("  [WARNING] Không thể xuất PDF. Cài thêm: docx2pdf | LibreOffice | pypandoc+pandoc")
    return None


# ── HTML Builder ─────────────────────────────────────────────────────────────

# Hex strings (reuse các hằng đã khai báo ở trên)
_H_NAVY_DARK   = "#251C53"
_H_NAVY_LIGHT  = "#4A3F8C"
_H_NAVY_PALE   = "#B8B3D6"
_H_ORANGE      = "#F05922"
_H_ORANGE_PALE = "#FFAA80"
_H_GRAY        = "#CBCBCB"
_H_WHITE       = "#FFFFFF"
_H_BG_CODE     = "#F4F4F8"


def _build_css() -> str:
    return f"""
/* ── Reset & Base ─────────────────────────────────────── */
*, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
html {{ font-size: 16px; scroll-behavior: smooth; }}
body {{
    font-family: 'Segoe UI', Arial, sans-serif;
    font-size: 10.5pt;
    color: #333;
    background: #F8F8FA;
    line-height: 1.65;
}}
/* ── Page wrapper ─────────────────────────────────────── */
.page {{
    max-width: 900px;
    margin: 0 auto;
    background: {_H_WHITE};
    box-shadow: 0 2px 16px rgba(37,28,83,.10);
    min-height: 100vh;
}}
/* ── Header ───────────────────────────────────────────── */
.kz-header {{
    display: flex;
    align-items: center;
    gap: 14px;
    padding: 14px 40px 12px;
    border-bottom: 3px solid {_H_ORANGE};
    background: {_H_WHITE};
    position: sticky;
    top: 0;
    z-index: 100;
}}
.kz-header img {{
    height: 36px;
    width: auto;
    object-fit: contain;
    flex-shrink: 0;
}}
.kz-header-name {{
    font-size: 9pt;
    font-weight: 700;
    color: {_H_NAVY_DARK};
    letter-spacing: .3px;
    text-transform: uppercase;
}}
/* ── Footer ───────────────────────────────────────────── */
.kz-footer {{
    text-align: center;
    padding: 16px 40px;
    font-size: 8.5pt;
    color: {_H_NAVY_LIGHT};
    border-top: 2px solid {_H_ORANGE};
    background: {_H_WHITE};
}}
.kz-footer a {{ color: {_H_NAVY_LIGHT}; text-decoration: none; }}
.kz-footer a:hover {{ color: {_H_ORANGE}; }}
/* ── Content ──────────────────────────────────────────── */
.content {{ padding: 32px 40px 40px; }}
/* ── Headings ─────────────────────────────────────────── */
h1 {{
    font-size: 22pt; color: {_H_NAVY_DARK}; font-weight: 700;
    margin: 16px 0 8px; padding-bottom: 6px;
    border-bottom: 2px solid {_H_NAVY_PALE}; line-height: 1.3;
}}
h2 {{
    font-size: 15pt; color: {_H_ORANGE}; font-weight: 700;
    margin: 22px 0 8px; padding-left: 10px;
    border-left: 4px solid {_H_ORANGE};
}}
h3 {{
    font-size: 12pt; color: {_H_NAVY_LIGHT}; font-weight: 700;
    margin: 16px 0 6px;
}}
h4 {{
    font-size: 11pt; color: {_H_NAVY_DARK}; font-weight: 700;
    margin: 12px 0 4px;
}}
h5, h6 {{
    font-size: 10.5pt; color: {_H_NAVY_DARK}; font-weight: 600;
    margin: 10px 0 4px;
}}
/* ── Paragraph ────────────────────────────────────────── */
p {{ margin: 0 0 10px; font-size: 10.5pt; }}
/* ── Links ────────────────────────────────────────────── */
a {{ color: {_H_NAVY_LIGHT}; text-decoration: underline; }}
a:hover {{ color: {_H_ORANGE}; }}
/* ── HR ───────────────────────────────────────────────── */
hr {{ border: none; border-top: 2px solid {_H_ORANGE}; margin: 20px 0; opacity: .5; }}
/* ── Blockquote ───────────────────────────────────────── */
blockquote {{
    border-left: 4px solid {_H_ORANGE}; margin: 12px 0;
    padding: 6px 16px; background: #FFF8F5;
    color: {_H_NAVY_LIGHT}; font-style: italic;
    border-radius: 0 4px 4px 0;
}}
blockquote p {{ margin: 0; }}
/* ── Code inline ──────────────────────────────────────── */
code {{
    font-family: 'Consolas', 'Courier New', monospace; font-size: 9pt;
    background: {_H_BG_CODE}; color: {_H_ORANGE};
    padding: 1px 5px; border-radius: 3px;
    border: 1px solid {_H_NAVY_PALE};
}}
/* ── Code block ───────────────────────────────────────── */
pre {{
    background: {_H_BG_CODE}; border: 1px solid {_H_NAVY_PALE};
    border-left: 3px solid {_H_NAVY_LIGHT}; border-radius: 4px;
    padding: 14px 16px; overflow-x: auto; margin: 12px 0;
    font-size: 9pt; line-height: 1.55;
}}
pre code {{ background: none; border: none; padding: 0; color: {_H_NAVY_DARK}; font-size: inherit; }}
/* ── Lists ────────────────────────────────────────────── */
ul, ol {{ margin: 6px 0 10px 24px; padding: 0; }}
li {{ margin-bottom: 3px; font-size: 10.5pt; }}
ul li::marker {{ color: {_H_ORANGE}; }}
ol li::marker {{ color: {_H_NAVY_DARK}; font-weight: 600; }}
/* ── Checkbox list ────────────────────────────────────── */
.checklist {{ list-style: none; margin-left: 4px; }}
.checklist li {{ display: flex; align-items: flex-start; gap: 8px; margin-bottom: 4px; }}
.checklist .cb {{ font-size: 13pt; line-height: 1.3; flex-shrink: 0; }}
.checklist .cb.checked {{ color: {_H_ORANGE}; }}
.checklist .cb.unchecked {{ color: {_H_GRAY}; }}
.checklist .checked-text {{ color: {_H_NAVY_LIGHT}; }}
/* ── Table ────────────────────────────────────────────── */
.table-wrap {{ overflow-x: auto; margin: 12px 0 16px; }}
table {{ width: 100%; border-collapse: collapse; font-size: 10pt; }}
th {{
    background: {_H_NAVY_DARK}; color: {_H_WHITE}; font-weight: 700;
    padding: 8px 12px; text-align: left; border: 1px solid {_H_NAVY_DARK};
}}
td {{ padding: 7px 12px; border: 1px solid {_H_GRAY}; vertical-align: top; }}
tr:nth-child(even) td {{ background: {_H_NAVY_PALE}26; }}
tr:hover td {{ background: {_H_ORANGE_PALE}22; }}
/* ── Images ───────────────────────────────────────────── */
.img-wrap {{ text-align: center; margin: 16px 0; }}
.img-wrap img {{
    max-width: 100%; border: 1px solid {_H_GRAY};
    border-radius: 4px; box-shadow: 0 2px 8px rgba(0,0,0,.08);
}}
.img-caption {{ font-size: 9pt; color: {_H_NAVY_LIGHT}; font-style: italic; margin-top: 6px; }}
.img-missing {{
    text-align: center; padding: 12px; background: #FFF3EE;
    border: 1px dashed {_H_ORANGE}; border-radius: 4px;
    color: {_H_ORANGE}; font-size: 9.5pt; margin: 12px 0;
}}
/* ── Print ────────────────────────────────────────────── */
@media print {{
    body {{ background: {_H_WHITE}; }}
    .page {{ box-shadow: none; max-width: 100%; }}
    .kz-header {{ position: static; }}
    pre {{ white-space: pre-wrap; word-break: break-word; }}
}}
/* ── Responsive ───────────────────────────────────────── */
@media (max-width: 640px) {{
    .content {{ padding: 20px 18px 28px; }}
    .kz-header {{ padding: 12px 18px; }}
    .kz-footer {{ padding: 12px 18px; }}
    h1 {{ font-size: 18pt; }}
    h2 {{ font-size: 13pt; }}
}}
"""


def _escape_html(text: str) -> str:
    return (text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;"))


def _parse_inline_html(text: str) -> str:
    """Chuyển đổi inline markdown → HTML (bold, italic, code, link, strikethrough)."""
    text = _escape_html(text)
    text = re.sub(r"\*\*\*(.+?)\*\*\*", r"<strong><em>\1</em></strong>", text)
    text = re.sub(r"\*\*(.+?)\*\*",     r"<strong>\1</strong>", text)
    text = re.sub(r"__(.+?)__",          r"<strong>\1</strong>", text)
    text = re.sub(r"\*([^*\n]+?)\*",    r"<em>\1</em>", text)
    text = re.sub(r"_([^_\n]+?)_",      r"<em>\1</em>", text)
    text = re.sub(r"~~(.+?)~~",          r"<del>\1</del>", text)
    text = re.sub(r"`([^`]+)`",          r"<code>\1</code>", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r'<a href="\2">\1</a>', text)
    return text


class _HtmlBuilder:
    """Parse Markdown và sinh HTML chunks."""

    def __init__(self, md_path: Path, embed_assets: bool = False):
        self.md_path = md_path
        self.embed_assets = embed_assets
        self.parts: list[str] = []

    def _flush_code(self, lang: str, lines: list[str]):
        code = _escape_html("\n".join(lines))
        lang_cls = f' class="language-{_escape_html(lang)}"' if lang else ""
        self.parts.append(f"<pre><code{lang_cls}>{code}</code></pre>\n")

    def _flush_table(self, rows: list[list[str]]):
        data = [r for r in rows if not re.match(r"^[\s\-|:]+$", "|".join(r))]
        if not data:
            return
        html = ['<div class="table-wrap"><table>']
        for ri, cells in enumerate(data):
            tag = "th" if ri == 0 else "td"
            html.append("<tr>")
            for cell in cells:
                html.append(f"<{tag}>{_parse_inline_html(cell.strip())}</{tag}>")
            html.append("</tr>")
        html.append("</table></div>\n")
        self.parts.append("\n".join(html))

    def _flush_checklist(self, items: list[tuple[bool, str, int]]):
        html = ['<ul class="checklist">']
        for checked, text, indent in items:
            sym_cls = "checked" if checked else "unchecked"
            sym = "☑" if checked else "☐"
            txt_cls = ' class="checked-text"' if checked else ""
            pad = f' style="padding-left:{indent * 16}px"' if indent else ""
            html.append(
                f'<li{pad}>'
                f'<span class="cb {sym_cls}">{sym}</span>'
                f'<span{txt_cls}>{_parse_inline_html(text)}</span>'
                f'</li>'
            )
        html.append("</ul>\n")
        self.parts.append("\n".join(html))

    def _resolve_img_src(self, src_raw: str) -> str | None:
        """Trả về src đã resolve (base64 hoặc absolute path), None nếu không tìm thấy."""
        resolved = self.md_path.parent / src_raw
        if resolved.exists():
            abs_path = str(resolved.resolve())
        elif os.path.exists(src_raw):
            abs_path = src_raw
        elif src_raw.startswith(("http", "data:")):
            return src_raw
        else:
            return None

        if self.embed_assets:
            uri = image_to_base64(abs_path)
            return uri if uri else abs_path
        return abs_path.replace("\\", "/")

    def parse(self, text: str):
        lines = text.splitlines()
        i = 0
        in_code = False; code_lang = ""; code_buf: list[str] = []
        in_table = False; table_buf: list[list[str]] = []
        in_checklist = False; checklist_buf: list[tuple[bool, str, int]] = []
        in_ul = False; ul_buf: list[str] = []
        in_ol = False; ol_buf: list[str] = []

        def close_ul():
            nonlocal in_ul, ul_buf
            if in_ul and ul_buf:
                self.parts.append("<ul>\n" + "".join(ul_buf) + "</ul>\n")
            in_ul = False; ul_buf = []

        def close_ol():
            nonlocal in_ol, ol_buf
            if in_ol and ol_buf:
                self.parts.append("<ol>\n" + "".join(ol_buf) + "</ol>\n")
            in_ol = False; ol_buf = []

        def close_checklist():
            nonlocal in_checklist, checklist_buf
            if in_checklist and checklist_buf:
                self._flush_checklist(checklist_buf)
            in_checklist = False; checklist_buf = []

        def close_table():
            nonlocal in_table, table_buf
            if in_table and table_buf:
                self._flush_table(table_buf)
            in_table = False; table_buf = []

        while i < len(lines):
            line = lines[i]

            # Code block
            if line.startswith("```"):
                close_ul(); close_ol(); close_checklist(); close_table()
                if not in_code:
                    in_code = True; code_lang = line[3:].strip(); code_buf = []
                else:
                    in_code = False; self._flush_code(code_lang, code_buf)
                    code_buf = []; code_lang = ""
                i += 1; continue

            if in_code:
                code_buf.append(line); i += 1; continue

            # Table
            if line.startswith("|"):
                close_ul(); close_ol(); close_checklist()
                if not in_table:
                    in_table = True; table_buf = []
                cells = [c.strip() for c in line.strip().strip("|").split("|")]
                table_buf.append(cells)
                i += 1; continue
            elif in_table:
                close_table()

            # Heading
            m = re.match(r"^(#{1,6})\s+(.*)", line)
            if m:
                close_ul(); close_ol(); close_checklist()
                level = len(m.group(1))
                content = _parse_inline_html(m.group(2).strip())
                anchor = re.sub(r"[^\w\- ]", "", m.group(2).strip().lower()).replace(" ", "-")
                self.parts.append(f'<h{level} id="{anchor}">{content}</h{level}>\n')
                i += 1; continue

            # HR
            if re.match(r"^[-*_]{3,}\s*$", line):
                close_ul(); close_ol(); close_checklist()
                self.parts.append("<hr>\n"); i += 1; continue

            # Blockquote
            if line.startswith("> "):
                close_ul(); close_ol(); close_checklist()
                content = _parse_inline_html(line[2:].strip())
                self.parts.append(f"<blockquote><p>{content}</p></blockquote>\n")
                i += 1; continue

            # Page break
            if "<!-- pagebreak -->" in line.lower():
                close_ul(); close_ol(); close_checklist()
                self.parts.append('<div style="page-break-after:always"></div>\n')
                i += 1; continue

            # Checkbox list
            if re.match(r"^(\s*)[-*] \[[ xX]\]", line):
                close_ul(); close_ol()
                m2 = re.match(r"^(\s*)[-*] \[([ xX])\] (.*)", line)
                checked = m2.group(2).lower() == "x"
                indent = len(m2.group(1)) // 2
                if not in_checklist:
                    in_checklist = True; checklist_buf = []
                checklist_buf.append((checked, m2.group(3).strip(), indent))
                i += 1; continue
            elif in_checklist:
                close_checklist()

            # Unordered list
            if re.match(r"^(\s*)[-*+] ", line):
                close_ol(); close_checklist()
                m2 = re.match(r"^(\s*)[-*+] (.*)", line)
                indent = len(m2.group(1)) // 2
                content = _parse_inline_html(m2.group(2).strip())
                if not in_ul:
                    in_ul = True; ul_buf = []
                pad = f' style="padding-left:{indent * 16}px"' if indent else ""
                ul_buf.append(f"  <li{pad}>{content}</li>\n")
                i += 1; continue
            elif in_ul:
                close_ul()

            # Ordered list
            if re.match(r"^(\s*)\d+\. ", line):
                close_ul(); close_checklist()
                m2 = re.match(r"^(\s*)\d+\. (.*)", line)
                indent = len(m2.group(1)) // 2
                content = _parse_inline_html(m2.group(2).strip())
                if not in_ol:
                    in_ol = True; ol_buf = []
                pad = f' style="padding-left:{indent * 16}px"' if indent else ""
                ol_buf.append(f"  <li{pad}>{content}</li>\n")
                i += 1; continue
            elif in_ol:
                close_ol()

            # Image (standalone line)
            img_m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line.strip())
            if img_m:
                close_ul(); close_ol(); close_checklist()
                alt = _escape_html(img_m.group(1))
                src = self._resolve_img_src(img_m.group(2))
                if src:
                    caption = f'<p class="img-caption">{alt}</p>' if alt else ""
                    self.parts.append(
                        f'<div class="img-wrap">'
                        f'<img src="{src}" alt="{alt}">'
                        f'{caption}</div>\n'
                    )
                else:
                    self.parts.append(
                        f'<div class="img-missing">⚠ ẢNH CHƯA CÓ: '
                        f'{_escape_html(img_m.group(2))}</div>\n'
                    )
                i += 1; continue

            # Blank line
            if line.strip() == "":
                close_ul(); close_ol(); close_checklist()
                i += 1; continue

            # Plain paragraph
            close_ul(); close_ol(); close_checklist()
            self.parts.append(f"<p>{_parse_inline_html(line.strip())}</p>\n")
            i += 1

        close_ul(); close_ol(); close_checklist(); close_table()
        if in_code and code_buf:
            self._flush_code(code_lang, code_buf)

    def get_html(self) -> str:
        return "".join(self.parts)


def _build_html_page(content_html: str, title: str,
                     logo_path: str | None, embed_assets: bool) -> str:
    if logo_path:
        if embed_assets:
            logo_src, _ = logo_to_base64(logo_path)
        else:
            logo_src = logo_path.replace("\\", "/")
        logo_tag = f'<img src="{logo_src}" alt="KZTEK Logo">'
    else:
        logo_tag = ""

    return f"""<!DOCTYPE html>
<html lang="vi">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{_escape_html(title)} — KZTEK</title>
  <style>{_build_css()}</style>
</head>
<body>
  <div class="page">
    <header class="kz-header">
      {logo_tag}
      <span class="kz-header-name">Công ty Cổ phần Đầu tư và Phát triển KZTEK</span>
    </header>
    <main class="content">
{content_html}
    </main>
    <footer class="kz-footer">
      <a href="https://kztek.net">kztek.net</a> &nbsp;|&nbsp;
      <a href="mailto:sales@kztek.net">sales@kztek.net</a> &nbsp;|&nbsp;
      0988&nbsp;637&nbsp;099
    </footer>
  </div>
</body>
</html>
"""


def convert_md_to_html(md_path: Path, output_path: Path,
                       logo_path: str | None, embed_assets: bool) -> Path:
    text = md_path.read_text(encoding="utf-8")
    title_m = re.search(r"^#\s+(.+)", text, re.MULTILINE)
    title = title_m.group(1).strip() if title_m else md_path.stem

    builder = _HtmlBuilder(md_path, embed_assets)
    builder.parse(text)
    html = _build_html_page(builder.get_html(), title, logo_path, embed_assets)
    output_path.write_text(html, encoding="utf-8")
    return output_path


# ── Main ─────────────────────────────────────────────────────────────────────

def process_file(
    md_path: Path,
    output_dir: Path | None,
    skip_docx: bool,
    skip_pdf: bool,
    logo: str | None,
    html: bool = False,
    embed_assets: bool = False,
    no_open: bool = False,
):
    """Xử lý 1 file Markdown → DOCX, PDF, và/hoặc HTML."""
    out_dir = output_dir or md_path.parent
    out_dir.mkdir(parents=True, exist_ok=True)

    print(f"\n→ Đang xử lý: {md_path}")

    docx_path = None
    if not skip_docx:
        docx_path = out_dir / md_path.with_suffix(".docx").name
        print(f"  Xuất DOCX : {docx_path}")
        convert_md_to_docx(md_path, docx_path, logo)
        print(f"  ✓ DOCX hoàn thành")

    if not skip_pdf:
        if docx_path and docx_path.exists():
            pdf_result = convert_docx_to_pdf(docx_path)
            if pdf_result:
                print(f"  ✓ PDF  hoàn thành: {pdf_result}")
            else:
                print(f"  ✗ PDF  thất bại (xem hướng dẫn ở trên)")
        else:
            print(f"  [SKIP] PDF bỏ qua vì không có DOCX (dùng --no-docx --no-pdf hoặc bỏ --no-docx)")

    if html:
        html_path = out_dir / md_path.with_suffix(".html").name
        print(f"  Xuất HTML : {html_path}")
        convert_md_to_html(md_path, html_path, logo, embed_assets)
        print(f"  ✓ HTML hoàn thành")
        if not no_open:
            try:
                webbrowser.open(html_path.resolve().as_uri())
            except Exception:
                pass

    return docx_path


def main():
    parser = argparse.ArgumentParser(
        description="Chuyển đổi Markdown → DOCX + PDF + HTML theo chuẩn thương hiệu KZTEK",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument(
        "inputs", nargs="+",
        help="File .md hoặc thư mục (dùng với --batch)"
    )
    parser.add_argument(
        "--output-dir", "-o", type=Path, default=None,
        help="Thư mục lưu file output (mặc định: cùng thư mục với file input)"
    )
    parser.add_argument(
        "--batch", "-b", action="store_true",
        help="Khi input là thư mục: chuyển toàn bộ file *.md bên trong"
    )
    parser.add_argument(
        "--no-docx", action="store_true",
        help="Bỏ qua xuất DOCX"
    )
    parser.add_argument(
        "--no-pdf", action="store_true",
        help="Bỏ qua xuất PDF"
    )
    parser.add_argument(
        "--html", action="store_true",
        help="Xuất thêm file HTML theo brand KZTEK"
    )
    parser.add_argument(
        "--embed-assets", "-e", action="store_true",
        help="Nhúng logo và ảnh dưới dạng base64 vào HTML (chỉ có tác dụng khi dùng --html)"
    )
    parser.add_argument(
        "--no-open", action="store_true",
        help="Không tự mở browser sau khi xuất HTML"
    )
    args = parser.parse_args()

    logo = find_logo()
    if logo:
        print(f"[KZTEK] Dùng logo: {logo}")
    else:
        print("[KZTEK] Không tìm thấy file logo — tiếp tục không có logo")

    md_files: list[Path] = []

    for inp in args.inputs:
        p = Path(inp)
        if p.is_dir():
            if args.batch:
                md_files.extend(sorted(p.glob("**/*.md")))
            else:
                parser.error(f"{p} là thư mục — thêm --batch để xử lý toàn bộ file bên trong")
        elif p.is_file() and p.suffix.lower() == ".md":
            md_files.append(p)
        else:
            print(f"[SKIP] Bỏ qua (không phải .md hoặc không tồn tại): {p}")

    if not md_files:
        sys.exit("[KZTEK] Không có file .md nào để xử lý.")

    print(f"\n[KZTEK] Sẽ xử lý {len(md_files)} file(s):")
    for f in md_files:
        print(f"  • {f}")

    success, failed = 0, 0
    for md_file in md_files:
        try:
            process_file(
                md_file,
                args.output_dir,
                skip_docx=args.no_docx,
                skip_pdf=args.no_pdf,
                logo=logo,
                html=args.html,
                embed_assets=args.embed_assets,
                no_open=args.no_open,
            )
            success += 1
        except Exception as e:
            print(f"  ✗ Lỗi khi xử lý {md_file}: {e}")
            failed += 1

    print(f"\n[KZTEK] Hoàn thành: {success} thành công, {failed} thất bại")


if __name__ == "__main__":
    main()
